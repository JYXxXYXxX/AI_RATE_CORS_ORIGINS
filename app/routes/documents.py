"""文档上传、分析、任务、报告相关路由。"""

from __future__ import annotations

import asyncio
from pathlib import Path
from typing import Any
from urllib.parse import quote

from fastapi import (
    APIRouter,
    BackgroundTasks,
    Depends,
    File,
    Form,
    HTTPException,
    Response,
    UploadFile,
)
from fastapi.responses import FileResponse

from app.config import Settings, get_settings
from app.db import get_repository
from app.pipeline import UnifiedPipeline
from app.reporting import render_report_markdown
from app.routes.deps import (
    AuthContext,
    ensure_document_access,
    get_auth_context,
    get_llm_rewrite_service,
    get_unified_pipeline,
    stage_from_status,
)
from app.plagiarism.scoring import score_duplication
from app.schemas_unified import (
    AnalysisRunStatusResponse,
    AnalysisTaskCreateResponse,
    AnalysisTaskStatusResponse,
    AnalyzeDocumentResponse,
    AttachReportResponse,
    CnkiReportSummary,
    CnkiRiskSpanResponse,
    DocumentBlockResponse,
    DocumentPatchRequest,
    DocumentPatchResponse,
    DocumentUploadResponse,
    ExportRequest,
    InternalRiskData,
    ManualReportSpanBindRequest,
    ManualReportSpanBindResponse,
    ReanalyzeRequest,
    ReanalyzeResponse,
    ReanalyzeSectionResult,
    ReportRiskData,
    RewriteAdviceResponse,
    UnifiedReportResponse,
    UploadWithReportResponse,
)
from app.services.analyzer import PaperAnalyzer, risk_level
from app.services.block_matcher import match_spans_to_blocks
from app.services.cnki_report_parser import parse_cnki_report_bytes
from app.services.document_blocks import parse_document_to_blocks
from app.services.feedback_learning import (
    append_feedback_learning_sample,
    append_private_feedback_learning_sample,
    build_feedback_learning_sample,
    refresh_feedback_learning_skill,
)
from app.services.text_processing import preview_text, segment_document
from app.task_queue import dispatch_analysis_task

router = APIRouter(prefix="/v1", tags=["documents"])


@router.post("/documents/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    title: str | None = Form(default=None),
    subject: str | None = Form(default=None),
    degree_level: str | None = Form(default=None),
    pipeline: UnifiedPipeline = Depends(get_unified_pipeline),
    auth: AuthContext = Depends(get_auth_context),
) -> DocumentUploadResponse:
    try:
        result = await pipeline.upload_document(
            file=file,
            title=title,
            subject=subject,
            degree_level=degree_level,
            user_id=str(auth.user["id"]) if auth.user is not None else None,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    document = result.document
    if auth.user is not None:
        pipeline.repository.grant_document_access(
            user_id=str(auth.user["id"]), document_id=str(document["id"])
        )
    return DocumentUploadResponse(
        document_id=str(document["id"]),
        title=document.get("title"),
        filename=document["filename"],
        subject=document.get("subject"),
        degree_level=document.get("degree_level"),
        char_count=document["char_count"],
        status=document["status"],
        reused_existing=result.reused_existing,
        created_at=document["created_at"],
    )


@router.post("/documents/upload-with-report", response_model=UploadWithReportResponse)
async def upload_document_with_report(
    file: UploadFile = File(...),
    report_file: UploadFile = File(...),
    title: str | None = Form(default=None),
    subject: str | None = Form(default=None),
    degree_level: str | None = Form(default=None),
    learning_consent: bool = Form(default=False),
    learning_scope: str = Form(default="none"),
    settings: Settings = Depends(get_settings),
    auth: AuthContext = Depends(get_auth_context),
) -> UploadWithReportResponse:
    """上传论文 + 知网检测报告，启用报告驱动模式。"""
    repository = get_repository()

    # 1. 上传论文（复用 pipeline 逻辑）
    pipeline = get_unified_pipeline()
    try:
        result = await pipeline.upload_document(
            file=file,
            title=title,
            subject=subject,
            degree_level=degree_level,
            user_id=str(auth.user["id"]) if auth.user is not None else None,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    document = result.document
    document_id = str(document["id"])
    if auth.user is not None:
        repository.grant_document_access(
            user_id=str(auth.user["id"]), document_id=document_id
        )

    # 2. 保存并解析知网报告
    report_content = await report_file.read()
    if not report_content:
        raise HTTPException(status_code=400, detail="报告文件为空")

    report_path = pipeline._write_binary(
        settings.upload_storage_dir, report_file.filename or "report.pdf", report_content
    )

    try:
        cnki_report = parse_cnki_report_bytes(
            report_file.filename or "report.pdf", report_content
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=f"报告解析失败: {exc}") from exc

    # 3. 获取论文 blocks
    blocks = parse_document_to_blocks(
        document["original_file_path"], document.get("source_type", "upload")
    )
    if blocks:
        repository.insert_document_blocks(document_id, [b.to_dict() for b in blocks])

    # 4. 三级匹配
    block_list = blocks if blocks else []
    mappings, unmatched = match_spans_to_blocks(cnki_report.risky_spans, block_list)

    # 5. 保存报告和 spans
    db_report = repository.insert_cnki_report(
        document_id=document_id,
        report_type=cnki_report.report_type,
        filename=report_file.filename,
        raw_format=Path(report_file.filename or "").suffix.lower().lstrip("."),
        total_copy_ratio=cnki_report.total_copy_ratio,
        aigc_ratio=cnki_report.aigc_ratio,
        generated_at=cnki_report.generated_at,
        raw_data=cnki_report.raw_meta,
        status="mapped" if mappings else "parsed",
    )
    report_id = str(db_report["id"])

    if cnki_report.risky_spans:
        repository.insert_cnki_report_spans(
            report_id,
            [
                {
                    "span_id": s.span_id,
                    "text": s.text,
                    "risk_type": s.risk_type,
                    "risk_level": s.risk_level,
                    "similarity": s.similarity,
                    "aigc_score": s.aigc_score,
                    "matched_source": s.matched_source,
                    "page_number": s.page_number,
                    "raw_meta": s.raw_meta,
                }
                for s in cnki_report.risky_spans
            ],
        )

    # 6. 保存映射并更新 block report_risk
    for mapping in mappings:
        repository.insert_block_report_mapping(
            document_id=document_id,
            block_id=mapping.block_id,
            span_id=mapping.span_id,
            report_id=report_id,
            match_method=mapping.match_method,
            match_confidence=mapping.match_confidence,
            matched_text=mapping.matched_text,
        )
        # 找到对应的 span 获取风险详情
        span = next((s for s in cnki_report.risky_spans if s.span_id == mapping.span_id), None)
        if span:
            repository.update_block_report_risk(
                document_id=document_id,
                block_id=mapping.block_id,
                report_risk={
                    "source": "cnki",
                    "risk_type": span.risk_type,
                    "risk_level": span.risk_level,
                    "similarity": span.similarity,
                    "aigc_score": span.aigc_score,
                    "matched_source": span.matched_source,
                    "span_id": span.span_id,
                    "match_confidence": mapping.match_confidence,
                },
            )

    # 7. 创建 report mode 的 analysis_run
    db_run = repository.create_analysis_run(
        document_id=document_id,
        run_type="report_driven",
        mode="report",
    )
    run_id = str(db_run["id"])
    repository.mark_run_completed(run_id)
    repository.mark_document_status(document_id, "completed")

    _record_official_report_learning_if_allowed(
        repository=repository,
        settings=settings,
        document=document,
        report=db_report,
        spans=[
            {
                "span_id": s.span_id,
                "text": s.text,
                "risk_type": s.risk_type,
                "risk_level": s.risk_level,
                "similarity": s.similarity,
                "aigc_score": s.aigc_score,
                "matched_source": s.matched_source,
                "page_number": s.page_number,
            }
            for s in cnki_report.risky_spans
        ],
        predicted_run_id=run_id,
        learning_consent=learning_consent,
        learning_scope=learning_scope,
        learning_user_id=auth.user_id if auth.authenticated else None,
    )

    # 8. 组装响应
    db_blocks = repository.list_document_blocks(document_id)
    block_responses = [_build_block_response(b, None) for b in db_blocks]

    risk_counts = {"high": 0, "medium": 0, "low": 0}
    for span in cnki_report.risky_spans:
        risk_counts[span.risk_level] = risk_counts.get(span.risk_level, 0) + 1

    return UploadWithReportResponse(
        file_id=document_id,
        run_id=run_id,
        report_mode=True,
        report_summary=CnkiReportSummary(
            report_id=report_id,
            report_type=cnki_report.report_type,
            total_copy_ratio=cnki_report.total_copy_ratio,
            aigc_ratio=cnki_report.aigc_ratio,
            high_risk_count=risk_counts.get("high", 0),
            medium_risk_count=risk_counts.get("medium", 0),
            low_risk_count=risk_counts.get("low", 0),
            unmatched_count=len(unmatched),
        ),
        blocks=block_responses,
        unmatched_risk_spans=[
            CnkiRiskSpanResponse(
                span_id=s.span_id,
                text=s.text,
                risk_type=s.risk_type,
                risk_level=s.risk_level,
                similarity=s.similarity,
                aigc_score=s.aigc_score,
                matched_source=s.matched_source,
                page_number=s.page_number,
            )
            for s in unmatched
        ],
    )


@router.post("/documents/{document_id}/report", response_model=AttachReportResponse)
async def attach_report_to_document(
    document_id: str,
    report_file: UploadFile = File(...),
    learning_consent: bool = Form(default=False),
    learning_scope: str = Form(default="none"),
    settings: Settings = Depends(get_settings),
    auth: AuthContext = Depends(get_auth_context),
) -> AttachReportResponse:
    """为已有论文上传知网报告并执行匹配。"""
    repository = get_repository()
    document = repository.get_document(document_id)
    if document is None:
        raise HTTPException(status_code=404, detail="document not found")
    ensure_document_access(document_id=document_id, auth=auth, repository=repository)

    report_content = await report_file.read()
    if not report_content:
        raise HTTPException(status_code=400, detail="报告文件为空")

    try:
        cnki_report = parse_cnki_report_bytes(
            report_file.filename or "report.pdf", report_content
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=f"报告解析失败: {exc}") from exc

    # 获取 blocks
    blocks = repository.list_document_blocks(document_id)
    if not blocks:
        # 尝试从原始文件重新解析
        from app.services.document_blocks import parse_document_to_blocks
        block_objs = parse_document_to_blocks(
            document["original_file_path"], document.get("source_type", "upload")
        )
        if block_objs:
            repository.insert_document_blocks(document_id, [b.to_dict() for b in block_objs])
            blocks = repository.list_document_blocks(document_id)

    block_list = blocks if blocks else []
    mappings, unmatched = match_spans_to_blocks(cnki_report.risky_spans, block_list)

    # 清除旧映射
    repository.clear_block_report_risks(document_id)
    old_reports = repository.list_cnki_reports_by_document(document_id)
    for old in old_reports:
        repository.delete_block_report_mappings_by_report(str(old["id"]))

    # 保存新报告
    db_report = repository.insert_cnki_report(
        document_id=document_id,
        report_type=cnki_report.report_type,
        filename=report_file.filename,
        raw_format=Path(report_file.filename or "").suffix.lower().lstrip("."),
        total_copy_ratio=cnki_report.total_copy_ratio,
        aigc_ratio=cnki_report.aigc_ratio,
        generated_at=cnki_report.generated_at,
        raw_data=cnki_report.raw_meta,
        status="mapped" if mappings else "parsed",
    )
    report_id = str(db_report["id"])

    if cnki_report.risky_spans:
        repository.insert_cnki_report_spans(
            report_id,
            [
                {
                    "span_id": s.span_id,
                    "text": s.text,
                    "risk_type": s.risk_type,
                    "risk_level": s.risk_level,
                    "similarity": s.similarity,
                    "aigc_score": s.aigc_score,
                    "matched_source": s.matched_source,
                    "page_number": s.page_number,
                    "raw_meta": s.raw_meta,
                }
                for s in cnki_report.risky_spans
            ],
        )

    for mapping in mappings:
        repository.insert_block_report_mapping(
            document_id=document_id,
            block_id=mapping.block_id,
            span_id=mapping.span_id,
            report_id=report_id,
            match_method=mapping.match_method,
            match_confidence=mapping.match_confidence,
            matched_text=mapping.matched_text,
        )
        span = next((s for s in cnki_report.risky_spans if s.span_id == mapping.span_id), None)
        if span:
            repository.update_block_report_risk(
                document_id=document_id,
                block_id=mapping.block_id,
                report_risk={
                    "source": "cnki",
                    "risk_type": span.risk_type,
                    "risk_level": span.risk_level,
                    "similarity": span.similarity,
                    "aigc_score": span.aigc_score,
                    "matched_source": span.matched_source,
                    "span_id": span.span_id,
                    "match_confidence": mapping.match_confidence,
                },
            )

    latest_runs = repository.list_completed_runs(document_id, limit=1)
    predicted_run_id = str(latest_runs[0]["id"]) if latest_runs else None
    _record_official_report_learning_if_allowed(
        repository=repository,
        settings=settings,
        document=document,
        report=db_report,
        spans=[
            {
                "span_id": s.span_id,
                "text": s.text,
                "risk_type": s.risk_type,
                "risk_level": s.risk_level,
                "similarity": s.similarity,
                "aigc_score": s.aigc_score,
                "matched_source": s.matched_source,
                "page_number": s.page_number,
            }
            for s in cnki_report.risky_spans
        ],
        predicted_run_id=predicted_run_id,
        learning_consent=learning_consent,
        learning_scope=learning_scope,
        learning_user_id=auth.user_id if auth.authenticated else None,
    )

    return AttachReportResponse(
        report_id=report_id,
        mapped_count=len(mappings),
        unmatched_count=len(unmatched),
        unmatched_spans=[
            CnkiRiskSpanResponse(
                span_id=s.span_id,
                text=s.text,
                risk_type=s.risk_type,
                risk_level=s.risk_level,
                similarity=s.similarity,
                aigc_score=s.aigc_score,
                matched_source=s.matched_source,
                page_number=s.page_number,
            )
            for s in unmatched
        ],
    )


@router.post("/documents/{document_id}/remap-report")
async def remap_report(
    document_id: str,
    auth: AuthContext = Depends(get_auth_context),
) -> AttachReportResponse:
    """重新执行报告与 blocks 的匹配（blocks 可能已修改）。"""
    repository = get_repository()
    document = repository.get_document(document_id)
    if document is None:
        raise HTTPException(status_code=404, detail="document not found")
    ensure_document_access(document_id=document_id, auth=auth, repository=repository)

    reports = repository.list_cnki_reports_by_document(document_id)
    if not reports:
        raise HTTPException(status_code=404, detail="该文档没有关联的知网报告")

    latest_report = reports[0]
    report_id = str(latest_report["id"])
    spans = repository.list_cnki_report_spans(report_id)

    blocks = repository.list_document_blocks(document_id)
    if not blocks:
        raise HTTPException(status_code=400, detail="文档没有解析出 blocks，无法重新匹配")

    from app.services.cnki_report_parser import CnkiRiskSpan
    cnki_spans = [
        CnkiRiskSpan(
            span_id=s["span_id"],
            text=s["text"],
            risk_type=s["risk_type"],
            risk_level=s["risk_level"],
            similarity=s.get("similarity"),
            aigc_score=s.get("aigc_score"),
            matched_source=s.get("matched_source"),
            page_number=s.get("page_number"),
        )
        for s in spans
    ]

    mappings, unmatched = match_spans_to_blocks(cnki_spans, blocks)

    # 清除旧映射
    repository.clear_block_report_risks(document_id)
    repository.delete_block_report_mappings_by_report(report_id)

    for mapping in mappings:
        repository.insert_block_report_mapping(
            document_id=document_id,
            block_id=mapping.block_id,
            span_id=mapping.span_id,
            report_id=report_id,
            match_method=mapping.match_method,
            match_confidence=mapping.match_confidence,
            matched_text=mapping.matched_text,
        )
        span = next((s for s in cnki_spans if s.span_id == mapping.span_id), None)
        if span:
            repository.update_block_report_risk(
                document_id=document_id,
                block_id=mapping.block_id,
                report_risk={
                    "source": "cnki",
                    "risk_type": span.risk_type,
                    "risk_level": span.risk_level,
                    "similarity": span.similarity,
                    "aigc_score": span.aigc_score,
                    "matched_source": span.matched_source,
                    "span_id": span.span_id,
                    "match_confidence": mapping.match_confidence,
                },
            )

    return AttachReportResponse(
        report_id=report_id,
        mapped_count=len(mappings),
        unmatched_count=len(unmatched),
        unmatched_spans=[
            CnkiRiskSpanResponse(
                span_id=s.span_id,
                text=s.text,
                risk_type=s.risk_type,
                risk_level=s.risk_level,
                similarity=s.similarity,
                aigc_score=s.aigc_score,
                matched_source=s.matched_source,
                page_number=s.page_number,
            )
            for s in unmatched
        ],
    )


@router.post("/documents/{document_id}/analyze", response_model=AnalyzeDocumentResponse)
async def analyze_uploaded_document(
    document_id: str,
    force: bool = False,
    pipeline: UnifiedPipeline = Depends(get_unified_pipeline),
    auth: AuthContext = Depends(get_auth_context),
) -> AnalyzeDocumentResponse:
    ensure_document_access(
        document_id=document_id, auth=auth, repository=get_repository()
    )
    try:
        result = await asyncio.to_thread(pipeline.analyze_document, document_id, force)
    except ValueError as exc:
        message = str(exc)
        status_code = 404 if "not found" in message else 400
        raise HTTPException(status_code=status_code, detail=message) from exc

    run = result["run"]
    return AnalyzeDocumentResponse(
        run_id=str(run["id"]),
        document_id=str(run["document_id"]),
        status=run["status"],
        created_at=run["created_at"],
        finished_at=run.get("finished_at"),
    )


@router.post(
    "/documents/{document_id}/analyze-async", response_model=AnalysisTaskCreateResponse
)
def analyze_uploaded_document_async(
    document_id: str,
    background_tasks: BackgroundTasks,
    auth: AuthContext = Depends(get_auth_context),
    settings: Settings = Depends(get_settings),
) -> AnalysisTaskCreateResponse:
    repository = get_repository()
    ensure_document_access(document_id=document_id, auth=auth, repository=repository)
    user_id = str(auth.user["id"]) if auth.user is not None else None
    task = repository.create_analysis_task(user_id=user_id, document_id=document_id)
    try:
        dispatch_analysis_task(
            settings=settings,
            task_id=str(task["id"]),
            document_id=document_id,
            user_id=user_id,
            credit_cost=0,
            background_tasks=background_tasks,
        )
    except Exception as exc:  # noqa: BLE001
        repository.mark_analysis_task_failed(str(task["id"]), error_message=str(exc))
        raise HTTPException(
            status_code=500, detail="failed to dispatch async task"
        ) from exc
    return AnalysisTaskCreateResponse(
        task_id=str(task["id"]),
        document_id=str(task["document_id"]),
        status=task["status"],
        progress=int(task["progress"]),
        created_at=task["created_at"],
    )


@router.get("/runs/{run_id}", response_model=AnalysisRunStatusResponse)
def get_analysis_run(
    run_id: str,
    pipeline: UnifiedPipeline = Depends(get_unified_pipeline),
    auth: AuthContext = Depends(get_auth_context),
) -> AnalysisRunStatusResponse:
    run = get_repository().get_run(run_id)
    if run is None:
        raise HTTPException(status_code=404, detail="run not found")
    ensure_document_access(
        document_id=str(run["document_id"]), auth=auth, repository=get_repository()
    )
    try:
        return pipeline.build_run_status(run_id)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc


@router.get("/runs/{run_id}/report", response_model=UnifiedReportResponse)
def get_unified_report(
    run_id: str,
    pipeline: UnifiedPipeline = Depends(get_unified_pipeline),
    auth: AuthContext = Depends(get_auth_context),
) -> UnifiedReportResponse:
    run = get_repository().get_run(run_id)
    if run is None:
        raise HTTPException(status_code=404, detail="run not found")
    ensure_document_access(
        document_id=str(run["document_id"]), auth=auth, repository=get_repository()
    )
    report = pipeline.get_report(run_id)
    if report is None:
        raise HTTPException(status_code=404, detail="report not found")
    return UnifiedReportResponse.model_validate(report)


@router.get("/runs/{run_id}/report/markdown")
def get_unified_report_markdown(
    run_id: str,
    pipeline: UnifiedPipeline = Depends(get_unified_pipeline),
    auth: AuthContext = Depends(get_auth_context),
) -> Response:
    run = get_repository().get_run(run_id)
    if run is None:
        raise HTTPException(status_code=404, detail="run not found")
    ensure_document_access(
        document_id=str(run["document_id"]), auth=auth, repository=get_repository()
    )
    report = pipeline.get_report(run_id)
    if report is None:
        raise HTTPException(status_code=404, detail="report not found")
    filename = f"{_safe_download_name(report.get('title') or run_id)}-report.md"
    markdown = render_report_markdown(report)
    return Response(
        content=markdown,
        media_type="text/markdown; charset=utf-8",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"
        },
    )


@router.get("/tasks/{task_id}", response_model=AnalysisTaskStatusResponse)
def get_analysis_task_status(
    task_id: str,
    auth: AuthContext = Depends(get_auth_context),
) -> AnalysisTaskStatusResponse:
    repository = get_repository()
    task = repository.get_analysis_task(task_id)
    if task is None:
        raise HTTPException(status_code=404, detail="task not found")
    ensure_document_access(
        document_id=str(task["document_id"]), auth=auth, repository=repository
    )
    return AnalysisTaskStatusResponse(
        task_id=str(task["id"]),
        document_id=str(task["document_id"]),
        run_id=str(task["run_id"]) if task.get("run_id") else None,
        title=task.get("title"),
        filename=task.get("filename"),
        status=task["status"],
        stage=stage_from_status(task["status"]),
        progress=int(task.get("progress", 0)),
        created_at=task["created_at"],
        started_at=task.get("started_at"),
        finished_at=task.get("finished_at"),
        error_message=task.get("error_message"),
    )


@router.get("/runs/{run_id}/sections")
def get_run_sections(
    run_id: str,
    auth: AuthContext = Depends(get_auth_context),
) -> list[dict[str, Any]]:
    """获取论文全文段落（排除参考文献、目录、致谢），附带分数信息。"""
    repository = get_repository()
    run = repository.get_run(run_id)
    if run is None:
        raise HTTPException(status_code=404, detail="run not found")
    ensure_document_access(
        document_id=str(run["document_id"]), auth=auth, repository=repository
    )

    # 获取所有段落
    sections = repository.list_document_sections(str(run["document_id"]))

    # 获取分数
    scores = repository.list_section_scores(run_id)
    score_map: dict[int, dict[str, Any]] = {}
    for sc in scores:
        idx = sc.get("section_index")
        if idx is not None:
            score_map.setdefault(idx, {})[sc.get("score_type", "")] = sc

    # 组装结果，排除 references / acknowledgement
    skip_types = {"references", "acknowledgement"}
    result: list[dict[str, Any]] = []
    for sec in sections:
        if sec.get("section_type") in skip_types:
            continue
        idx = sec.get("section_index")
        aigc = score_map.get(idx, {}).get("aigc", {})
        dup = score_map.get(idx, {}).get("duplication", {})
        result.append(
            {
                "section_index": idx,
                "paragraph_index": sec.get("paragraph_index"),
                "section_title": sec.get("section_title"),
                "section_type": sec.get("section_type"),
                "content": sec.get("content") or sec.get("text_preview") or "",
                "char_count": sec.get("char_count", 0),
                "aigc_score": float(aigc.get("normalized_score", 0)) if aigc else 0.0,
                "dup_score": float(dup.get("normalized_score", 0)) if dup else 0.0,
                "risk_level": aigc.get("risk_level", "low") if aigc else "low",
                "reasons": aigc.get("reasons", []) if aigc else [],
            }
        )

    return result


def _safe_download_name(value: str) -> str:
    cleaned = "".join(
        char if char.isascii() and (char.isalnum() or char in {"-", "_"}) else "_"
        for char in value
    )
    return cleaned.strip("_") or "report"


@router.post("/runs/{run_id}/reanalyze", response_model=ReanalyzeResponse)
async def reanalyze_run(
    run_id: str,
    payload: ReanalyzeRequest,
    auth: AuthContext = Depends(get_auth_context),
) -> ReanalyzeResponse:
    """接收改写后的段落，重新计算 AIGC 和查重分数，保持 section_index 不变。"""
    repository = get_repository()
    run = repository.get_run(run_id)
    if run is None:
        raise HTTPException(status_code=404, detail="run not found")
    ensure_document_access(
        document_id=str(run["document_id"]), auth=auth, repository=repository
    )
    if run.get("mode") == "report":
        raise HTTPException(
            status_code=409,
            detail="该任务已上传官方查重/AIGC报告，复核结果应以官方报告为准。请上传新的官方复检报告来更新风险颜色和指标。",
        )

    original_sections = repository.list_document_sections(str(run["document_id"]))
    rewritten_map = {s.section_index: s.content for s in payload.sections}

    # 按原始顺序组装当前文本（优先使用改写后内容）
    text_parts: list[str] = []
    for sec in sorted(original_sections, key=lambda s: s.get("section_index", 0)):
        idx = sec.get("section_index", 0)
        text = rewritten_map.get(idx, sec.get("content") or sec.get("text_preview") or "")
        text_parts.append(text)
    full_text = "\n\n".join(text_parts)

    if not full_text.strip():
        raise HTTPException(status_code=400, detail="text is empty after assembly")

    settings = get_settings()

    # 为了保持 section_index 一致，直接对输入的每个 section 评分，而不是重新 segment
    from app.services.text_processing import TextSegment

    input_segments: list[TextSegment] = []
    for sec in payload.sections:
        input_segments.append(
            TextSegment(
                index=sec.section_index,
                text=sec.content,
                section_title=None,
                paragraph_index=sec.section_index,
            )
        )

    # AIGC 重算（在线程池中执行，避免阻塞事件循环）
    analyzer = PaperAnalyzer(settings)
    scored_reports = await asyncio.to_thread(
        lambda: [analyzer._score_segment(seg, input_segments) for seg in input_segments]
    )

    total_chars = sum(r.char_count for r in scored_reports)
    scored_chars = total_chars
    if scored_chars == 0:
        ai_like_score = 0.0
    else:
        weighted = sum(r.ai_like_score * r.char_count for r in scored_reports)
        ai_like_score = weighted / scored_chars

    dispersion = PaperAnalyzer._segment_dispersion(scored_reports)
    predicted_range = analyzer.calibrator.predict_range(
        ai_like_score, dispersion, run.get("subject")
    )
    confidence = analyzer.calibrator.confidence(len(scored_reports), dispersion)

    # 查重重算（本地段落间相似度 + 模板检测，无需 embedding）
    dup_sections: list[dict[str, Any]] = []
    for sec in payload.sections:
        dup_sections.append(
            {
                "section_index": sec.section_index,
                "content": sec.content,
                "text_preview": preview_text(sec.content),
                "char_count": len(sec.content),
                "section_title": None,
            }
        )
    duplication = await asyncio.to_thread(score_duplication, dup_sections)

    dup_map = {s.section_index: s for s in duplication.section_scores}
    sections_result: list[ReanalyzeSectionResult] = []
    for rep in scored_reports:
        dup_sec = dup_map.get(rep.index)
        dup_score = dup_sec.normalized_score if dup_sec else 0.0
        combined_risk = risk_level(max(rep.ai_like_score, dup_score))
        sections_result.append(
            ReanalyzeSectionResult(
                section_index=rep.index,
                section_title=rep.section_title,
                aigc_score=rep.ai_like_score,
                duplication_score=dup_score,
                risk_level=combined_risk,
            )
        )

    return ReanalyzeResponse(
        ai_like_score=round(ai_like_score, 4),
        ai_like_percent=round(ai_like_score * 100, 2),
        duplication_score=duplication.overall_score,
        duplication_percent=round(duplication.overall_score * 100, 2),
        risk_level=risk_level(max(ai_like_score, duplication.overall_score)),
        predicted_cnki_range=f"{predicted_range.label} ({predicted_range.lower_percent:.1f}%-{predicted_range.upper_percent:.1f}%)",
        confidence=f"{confidence:.2f}",
        segment_count=len(scored_reports),
        total_chars=total_chars,
        sections=sections_result,
        disclaimer="本报告为改写后重新计算的 AIGC 疑似风险与知网结果区间预测，不等同于知网、维普、万方或 Turnitin 官方结果。",
    )


@router.get("/documents/{document_id}/original")
async def download_original_document(
    document_id: str,
    auth: AuthContext = Depends(get_auth_context),
):
    """下载用户上传的原始文档文件（docx/pdf/txt等）。"""
    repository = get_repository()
    document = repository.get_document(document_id)
    if document is None:
        raise HTTPException(status_code=404, detail="document not found")
    ensure_document_access(
        document_id=document_id, auth=auth, repository=repository
    )

    original_path = document.get("original_file_path")
    if not original_path:
        raise HTTPException(status_code=404, detail="original file not found")

    file_path = Path(original_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="original file missing on disk")

    suffix = file_path.suffix.lower()
    media_type_map = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".pdf": "application/pdf",
        ".txt": "text/plain; charset=utf-8",
        ".md": "text/plain; charset=utf-8",
    }
    media_type = media_type_map.get(suffix, "application/octet-stream")
    filename = document.get("filename") or f"original{suffix}"

    return FileResponse(
        path=str(file_path),
        media_type=media_type,
        filename=filename,
    )


@router.post("/runs/{run_id}/export")
async def export_rewritten_document(
    run_id: str,
    payload: ExportRequest,
    auth: AuthContext = Depends(get_auth_context),
):
    """导出改写后的文档（.docx 或 .txt）。"""
    repository = get_repository()
    db_run = repository.get_run(run_id)
    if db_run is None:
        raise HTTPException(status_code=404, detail="run not found")
    ensure_document_access(
        document_id=str(db_run["document_id"]), auth=auth, repository=repository
    )

    # 付费功能已隐藏：导出始终开放
    # if auth.user:
    #     unlock = repository.get_run_unlock(str(auth.user["id"]), run_id, "export_docx")
    #     if not unlock or unlock.get("status") != "unlocked":
    #         raise HTTPException(status_code=402, detail="export_docx not unlocked")

    original_sections = repository.list_document_sections(str(db_run["document_id"]))
    rewritten_map = {s.section_index: s.content for s in payload.sections}

    # 组装全文，保持原有顺序
    assembled: list[tuple[int, str | None, str]] = []
    for sec in sorted(original_sections, key=lambda s: s.get("section_index", 0)):
        idx = sec.get("section_index", 0)
        title = sec.get("section_title")
        text = rewritten_map.get(idx, sec.get("content") or sec.get("text_preview") or "")
        assembled.append((idx, title, text))

    doc_title = db_run.get("title") or "rewritten"
    safe_title = _safe_download_name(doc_title) or "rewritten"

    if payload.format == "txt":
        if not assembled:
            try:
                blocks_for_txt = repository.list_document_blocks(str(db_run["document_id"]))
                patches_for_txt = repository.list_latest_patches_by_run(
                    str(db_run["document_id"]), run_id
                )
                patch_map_for_txt = {
                    patch["block_id"]: patch for patch in patches_for_txt
                }
                for block in sorted(
                    blocks_for_txt, key=lambda item: item.get("display_order", 0)
                ):
                    if block.get("block_type") in {"title", "heading"}:
                        continue
                    patch = patch_map_for_txt.get(block.get("block_id"))
                    text = (
                        patch.get("new_text")
                        if patch
                        else block.get("text") or ""
                    )
                    if text.strip():
                        assembled.append(
                            (
                                int(block.get("display_order", 0)),
                                block.get("section_title"),
                                text,
                            )
                        )
            except Exception:
                assembled = []
        lines: list[str] = []
        if doc_title:
            lines.append(doc_title)
            lines.append("=" * 40)
            lines.append("")
        for _idx, sec_title, text in assembled:
            if sec_title:
                lines.append(sec_title)
                lines.append("-" * 20)
            lines.append(text)
            lines.append("")
        content = "\n".join(lines)
        return Response(
            content=content,
            media_type="text/plain; charset=utf-8",
            headers={
                "Content-Disposition": f'attachment; filename="{safe_title}.txt"'
            },
        )

    # docx
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml
    except ImportError as exc:
        raise HTTPException(status_code=501, detail="docx export not available") from exc

    # 风险颜色映射（浅色背景，适合 Word 文档）
    RISK_COLORS = {
        "high": "FFCDD2",
        "medium": "FFE0B2",
        "low": "E1BEE7",
        "normal": "C8E6C9",
    }

    def _set_para_shading(paragraph, color_hex: str):
        pPr = paragraph._element.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:fill="{color_hex}" w:val="clear"/>'
        )
        pPr.append(shading)

    # 尝试基于原始 docx + blocks + patches 导出（新架构）
    original_path = None
    document = repository.get_document(str(db_run["document_id"]))
    if document:
        original_path = document.get("original_file_path")

    # 优先使用 blocks + patches 模式导出
    try:
        blocks = repository.list_document_blocks(str(db_run["document_id"]))
        patches = repository.list_latest_patches_by_run(
            str(db_run["document_id"]), run_id
        )
    except Exception:
        blocks = []
        patches = []

    if blocks and original_path and Path(original_path).suffix.lower() == ".docx" and Path(original_path).exists():
        try:
            from app.services.docx_patch import (
                DocxPatchError,
                export_docx_with_patch_report,
            )

            patch_result = export_docx_with_patch_report(
                original_path,
                blocks,
                patches,
                strict=True,
            )
            headers = {
                "Content-Disposition": f'attachment; filename="{safe_title}.docx"',
                **patch_result.stats.to_headers(),
            }
            return Response(
                content=patch_result.content,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers=headers,
            )
        except DocxPatchError as exc:
            raise HTTPException(
                status_code=409,
                detail={
                    "message": "DOCX 原文档文本替换失败。为避免破坏论文格式，系统未导出重建版文档。",
                    "stats": {
                        "requested_patch_count": exc.stats.requested_patch_count,
                        "applied_count": exc.stats.applied_count,
                        "failed_count": exc.stats.failed_count,
                        "skipped_block_count": exc.stats.skipped_block_count,
                    },
                    "failures": [
                        {
                            "block_id": item.block_id,
                            "paragraph_index": item.paragraph_index,
                            "reason": item.reason,
                            "old_text_preview": item.old_text_preview,
                        }
                        for item in exc.stats.failures[:10]
                    ],
                },
            ) from exc
        except Exception as exc:
            raise HTTPException(
                status_code=409,
                detail="DOCX 原文档导出失败。为避免破坏论文格式，系统未回退生成重建版文档。",
            ) from exc

    # 回退：基于 sections 的导出（兼容旧数据）
    risk_map = {s.section_index: s.risk_level for s in payload.sections}
    doc = Document()
    if doc_title:
        heading = doc.add_heading(doc_title, level=0)
        heading.alignment = 1  # center

    for idx, sec_title, text in assembled:
        if sec_title:
            doc.add_heading(sec_title, level=1)
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = "宋体"
        # 添加风险背景色
        risk = risk_map.get(idx, "normal")
        _set_para_shading(p, RISK_COLORS.get(risk, "C8E6C9"))

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_title}.docx"'
        },
    )


@router.post(
    "/runs/{run_id}/sections/{section_index}/rewrite-advice",
    response_model=RewriteAdviceResponse,
)
@router.get(
    "/runs/{run_id}/sections/{section_index}/rewrite-advice",
    response_model=RewriteAdviceResponse,
)
async def get_section_rewrite_advice(
    run_id: str,
    section_index: int,
    pipeline: UnifiedPipeline = Depends(get_unified_pipeline),
    llm_service=Depends(get_llm_rewrite_service),
    auth: AuthContext = Depends(get_auth_context),
) -> RewriteAdviceResponse:
    repository = get_repository()
    run = repository.get_run(run_id)
    if run is None:
        raise HTTPException(status_code=404, detail="run not found")
    ensure_document_access(
        document_id=str(run["document_id"]), auth=auth, repository=repository
    )

    # 付费功能已隐藏：改写始终开放
    # if auth.user:
    #     unlock = get_repository().get_run_unlock(str(auth.user["id"]), run_id, "unlock_rewrite")
    #     if not unlock or unlock.get("status") != "unlocked":
    #         raise HTTPException(status_code=402, detail="unlock_rewrite not unlocked")

    if run.get("mode") == "report":
        return await _get_report_mode_rewrite_advice(
            run=run,
            run_id=run_id,
            section_index=section_index,
            repository=repository,
            llm_service=llm_service,
        )

    report = pipeline.get_report(run_id)
    if report is None:
        raise HTTPException(status_code=404, detail="report not found")

    # 从报告中找到对应段落的风险信息
    section_info = None
    for item in report.get("top_risk_sections", []):
        if item.get("section_index") == section_index:
            section_info = item
            break
    if section_info is None:
        # 也尝试从 chapter_heatmap 里的 section 找，或者直接允许对任意段落改写
        for item in report.get("chapter_heatmap", []):
            # chapter_heatmap 没有 section_index，跳过精确匹配
            pass
        # 如果没有找到，也允许继续，只是缺少风险信息

    # 获取完整段落文本
    sections = repository.list_document_sections(str(run["document_id"]))
    target_section = None
    for sec in sections:
        if sec.get("section_index") == section_index:
            target_section = sec
            break
    if target_section is None:
        raise HTTPException(status_code=404, detail="section not found")

    text = target_section.get("content") or target_section.get("text_preview") or ""
    if not text.strip():
        raise HTTPException(status_code=400, detail="section text is empty")

    # 判断风险类型
    aigc_score = float(section_info.get("aigc_score", 0)) if section_info else 0.0
    dup_score = float(section_info.get("duplication_score", 0)) if section_info else 0.0
    if aigc_score >= 0.65 and aigc_score > dup_score:
        risk_type = "aigc"
    elif dup_score >= 0.50 and dup_score >= aigc_score:
        risk_type = "duplication"
    else:
        risk_type = "mixed"

    reasons = section_info.get("reasons") if section_info else []
    subject = report.get("subject")
    degree_level = report.get("degree_level")
    section_title = (
        target_section.get("section_title") or section_info.get("title")
        if section_info
        else None
    )

    # 获取最新知网实测数据，注入 prompt 让建议更精准
    latest_cnki_dup = None
    latest_cnki_aigc = None
    feedback_timeline = report.get("feedback_timeline", [])
    if feedback_timeline:
        latest = feedback_timeline[0]
        latest_cnki_dup = latest.get("cnki_dup_percent")
        latest_cnki_aigc = latest.get("cnki_aigc_percent")

    # 获取该段落的本地分数，用于计算知网-本地差距
    local_aigc_score = float(section_info.get("aigc_score", 0)) if section_info else 0.0
    local_dup_score = (
        float(section_info.get("duplication_score", 0)) if section_info else 0.0
    )

    result = await llm_service.rewrite_paragraph(
        text=text,
        risk_type=risk_type,
        reasons=reasons,
        subject=subject,
        section_title=section_title,
        degree_level=degree_level,
        cnki_dup_percent=latest_cnki_dup,
        cnki_aigc_percent=latest_cnki_aigc,
        local_aigc_score=local_aigc_score,
        local_dup_score=local_dup_score,
    )

    if result.get("error"):
        return RewriteAdviceResponse(
            run_id=run_id,
            section_index=section_index,
            diagnosis="",
            sentences=[],
            rewritten_paragraph="",
            overall_advice="",
            error=result["error"],
        )

    sentences = [
        {
            "original": s.get("original", ""),
            "risk": s.get("risk", "medium"),
            "rewritten": s.get("rewritten", ""),
            "explanation": s.get("explanation", ""),
        }
        for s in result.get("sentences", [])
    ]

    return RewriteAdviceResponse(
        run_id=run_id,
        section_index=section_index,
        diagnosis=result.get("diagnosis", ""),
        sentences=sentences,
        rewritten_paragraph=result.get("rewritten_paragraph", ""),
        overall_advice=result.get("overall_advice", ""),
    )


async def _get_report_mode_rewrite_advice(
    *,
    run: dict[str, Any],
    run_id: str,
    section_index: int,
    repository: Any,
    llm_service: Any,
) -> RewriteAdviceResponse:
    document_id = str(run["document_id"])
    blocks = repository.list_document_blocks(document_id)
    target_block = _find_report_mode_block(blocks, section_index)
    if target_block is None:
        raise HTTPException(status_code=404, detail="report block not found")

    report_risk = target_block.get("report_risk") or {}
    if not report_risk:
        raise HTTPException(
            status_code=400,
            detail="该段落未被官方查重/AIGC报告标记为风险段，报告模式下仅对官方标记片段生成改写建议。",
        )
    text = target_block.get("text") or ""
    if not text.strip():
        raise HTTPException(status_code=400, detail="block text is empty")

    latest_report = _latest_cnki_report(repository, document_id)
    cnki_dup = latest_report.get("total_copy_ratio") if latest_report else None
    cnki_aigc = latest_report.get("aigc_ratio") if latest_report else None

    risk_type = report_risk.get("risk_type") or "mixed"
    if risk_type not in {"aigc", "duplication", "similarity", "mixed"}:
        risk_type = "mixed"
    if risk_type == "similarity":
        risk_type = "duplication"

    risk_label = report_risk.get("risk_level", "medium")
    metric = report_risk.get("similarity")
    if metric is None:
        metric = report_risk.get("aigc_score")
    reasons = [
        f"官方报告标记为{risk_label}风险",
        f"官方报告类型：{report_risk.get('risk_type', 'mixed')}",
    ]
    if metric is not None:
        reasons.append(f"官方报告片段分值：{float(metric):.1f}%")
    if report_risk.get("matched_source"):
        reasons.append(f"相似来源：{report_risk['matched_source']}")

    result = await llm_service.rewrite_paragraph(
        text=text,
        risk_type=risk_type,
        reasons=reasons,
        subject=run.get("subject"),
        section_title=target_block.get("section_title"),
        degree_level=run.get("degree_level"),
        cnki_dup_percent=cnki_dup,
        cnki_aigc_percent=cnki_aigc,
        local_aigc_score=None,
        local_dup_score=None,
    )

    if result.get("error"):
        return RewriteAdviceResponse(
            run_id=run_id,
            section_index=section_index,
            diagnosis="",
            sentences=[],
            rewritten_paragraph="",
            overall_advice="",
            error=result["error"],
        )

    sentences = [
        {
            "original": s.get("original", ""),
            "risk": s.get("risk", "medium"),
            "rewritten": s.get("rewritten", ""),
            "explanation": s.get("explanation", ""),
        }
        for s in result.get("sentences", [])
    ]

    diagnosis = result.get("diagnosis", "")
    official_prefix = "本次改写以用户上传的官方查重/AIGC报告为准；风险等级、优先级和改写方向均围绕官方标记片段展开。"
    return RewriteAdviceResponse(
        run_id=run_id,
        section_index=section_index,
        diagnosis=f"{official_prefix}{diagnosis}",
        sentences=sentences,
        rewritten_paragraph=result.get("rewritten_paragraph", ""),
        overall_advice=result.get("overall_advice", ""),
    )


def _find_report_mode_block(
    blocks: list[dict[str, Any]], section_index: int
) -> dict[str, Any] | None:
    for block in blocks:
        source_map = block.get("source_map") or {}
        if source_map.get("paragraphIndex") == section_index:
            return block
    for block in blocks:
        if block.get("display_order") == section_index:
            return block
    return None


def _latest_cnki_report(
    repository: Any, document_id: str
) -> dict[str, Any] | None:
    reports = repository.list_cnki_reports_by_document(document_id)
    return reports[0] if reports else None


def _span_response(span: dict[str, Any]) -> CnkiRiskSpanResponse:
    return CnkiRiskSpanResponse(
        span_id=span["span_id"],
        text=span["text"],
        risk_type=span["risk_type"],
        risk_level=span["risk_level"],
        similarity=span.get("similarity"),
        aigc_score=span.get("aigc_score"),
        matched_source=span.get("matched_source"),
        page_number=span.get("page_number"),
    )


def _report_risk_from_span(
    span: dict[str, Any], *, match_confidence: float
) -> dict[str, Any]:
    return {
        "source": "cnki",
        "risk_type": span["risk_type"],
        "risk_level": span["risk_level"],
        "similarity": span.get("similarity"),
        "aigc_score": span.get("aigc_score"),
        "matched_source": span.get("matched_source"),
        "span_id": span["span_id"],
        "match_confidence": match_confidence,
    }


def _list_report_spans(repository: Any, report_id: str) -> list[dict[str, Any]]:
    if not hasattr(repository, "list_cnki_report_spans"):
        return []
    return repository.list_cnki_report_spans(report_id)


def _list_unmatched_report_spans(
    repository: Any, *, document_id: str, report_id: str
) -> list[dict[str, Any]]:
    if hasattr(repository, "list_unmapped_cnki_report_spans"):
        return repository.list_unmapped_cnki_report_spans(report_id)

    spans = _list_report_spans(repository, report_id)
    if not spans or not hasattr(repository, "list_block_report_mappings"):
        return []
    mapped_span_ids = {
        item.get("span_id")
        for item in repository.list_block_report_mappings(document_id)
        if str(item.get("report_id")) == str(report_id)
    }
    return [span for span in spans if span.get("span_id") not in mapped_span_ids]


def _record_official_report_learning_if_allowed(
    *,
    repository: Any,
    settings: Settings,
    document: dict[str, Any],
    report: dict[str, Any],
    spans: list[dict[str, Any]],
    predicted_run_id: str | None,
    learning_consent: bool,
    learning_scope: str = "none",
    learning_user_id: str | None = None,
) -> tuple[bool, bool]:
    effective_scope = _normalize_learning_scope(learning_scope, learning_consent)
    if effective_scope == "none":
        return False, False

    details = {
        "fragments": [
            {
                "type": span.get("risk_type") or "unknown",
                "source_text": span.get("text") or "",
                "similar_text": span.get("matched_source") or "",
                "matched_section_index": None,
                "match_ratio": span.get("similarity") or span.get("aigc_score"),
            }
            for span in spans[:80]
        ]
    }
    feedback_like = {
        "id": f"official-report:{report['id']}",
        "document_id": document["id"],
        "predicted_run_id": predicted_run_id,
        "cnki_dup_percent": report.get("total_copy_ratio"),
        "cnki_aigc_percent": report.get("aigc_ratio"),
        "report_date": report.get("generated_at"),
        "created_at": report.get("parsed_at"),
    }
    try:
        learning_sample = build_feedback_learning_sample(
            repository=repository,
            document=document,
            feedback=feedback_like,
            predicted_run_id=predicted_run_id,
            details=details,
        )
        if learning_sample:
            learning_sample["learning_scope"] = effective_scope
        if not learning_sample:
            return False, False
        if effective_scope == "anonymous_global":
            saved = append_feedback_learning_sample(
                settings.feedback_learning_store_path,
                learning_sample,
            )
        else:
            saved = append_private_feedback_learning_sample(
                settings.feedback_private_learning_dir,
                learning_user_id,
                learning_sample,
            )
        refreshed = (
            refresh_feedback_learning_skill(
                settings.feedback_learning_store_path,
                settings.feedback_learning_skill_path,
            )
            if saved and effective_scope == "anonymous_global"
            else False
        )
        return saved, refreshed
    except Exception:
        return False, False


def _normalize_learning_scope(scope: str | None, consent: bool = False) -> str:
    if scope == "private_account":
        return "private_account"
    if scope == "anonymous_global" or consent:
        return "anonymous_global"
    if scope == "none":
        return scope
    return "none"


# ------------------------------------------------------------------
# Document Blocks
# ------------------------------------------------------------------

@router.post(
    "/runs/{run_id}/report-spans/bind",
    response_model=ManualReportSpanBindResponse,
)
async def bind_report_span_to_block(
    run_id: str,
    payload: ManualReportSpanBindRequest,
    auth: AuthContext = Depends(get_auth_context),
) -> ManualReportSpanBindResponse:
    repository = get_repository()
    db_run = repository.get_run(run_id)
    if db_run is None:
        raise HTTPException(status_code=404, detail="run not found")

    document_id = str(db_run["document_id"])
    ensure_document_access(
        document_id=document_id, auth=auth, repository=repository
    )

    latest_report = _latest_cnki_report(repository, document_id)
    if latest_report is None:
        raise HTTPException(status_code=404, detail="official report not found")
    report_id = str(latest_report["id"])

    if hasattr(repository, "get_cnki_report_span"):
        span = repository.get_cnki_report_span(report_id, payload.span_id)
    else:
        span = next(
            (item for item in _list_report_spans(repository, report_id) if item.get("span_id") == payload.span_id),
            None,
        )
    if span is None:
        raise HTTPException(status_code=404, detail="report span not found")

    if hasattr(repository, "list_block_report_mappings"):
        existing_mappings = repository.list_block_report_mappings(document_id)
        for mapping in existing_mappings:
            if (
                str(mapping.get("report_id")) == report_id
                and mapping.get("span_id") == payload.span_id
            ):
                raise HTTPException(status_code=409, detail="report span already mapped")
    else:
        existing_mappings = []

    block = repository.get_document_block(document_id, payload.block_id)
    if block is None:
        raise HTTPException(status_code=404, detail="document block not found")

    repository.insert_block_report_mapping(
        document_id=document_id,
        block_id=payload.block_id,
        span_id=payload.span_id,
        report_id=report_id,
        match_method="manual",
        match_confidence=1.0,
        matched_text=block.get("text"),
    )
    updated_block = repository.update_block_report_risk(
        document_id=document_id,
        block_id=payload.block_id,
        report_risk=_report_risk_from_span(span, match_confidence=1.0),
    )

    unmatched = _list_unmatched_report_spans(
        repository, document_id=document_id, report_id=report_id
    )
    mapped_count = len(
        [
            item
            for item in repository.list_block_report_mappings(document_id)
            if str(item.get("report_id")) == report_id
        ]
    ) if hasattr(repository, "list_block_report_mappings") else len(existing_mappings) + 1

    return ManualReportSpanBindResponse(
        report_id=report_id,
        span_id=payload.span_id,
        block_id=payload.block_id,
        mapped_count=mapped_count,
        unmatched_count=len(unmatched),
        unmatched_spans=[_span_response(item) for item in unmatched],
        block=_build_block_response(updated_block or block, None),
    )


def _build_block_response(
    block: dict[str, Any], patch: dict[str, Any] | None
) -> DocumentBlockResponse:
    """辅助函数：从数据库行构建 DocumentBlockResponse。"""
    effective_text = patch["new_text"] if patch else block["text"]
    report_risk = None
    if block.get("report_risk"):
        rr = block["report_risk"]
        report_risk = ReportRiskData(
            source=rr.get("source", "cnki"),
            risk_type=rr.get("risk_type", "similarity"),
            risk_level=rr.get("risk_level", "medium"),
            similarity=rr.get("similarity"),
            aigc_score=rr.get("aigc_score"),
            matched_source=rr.get("matched_source"),
            span_id=rr.get("span_id", ""),
            match_confidence=rr.get("match_confidence", 0.0),
        )
    internal_risk = None
    if block.get("internal_risk"):
        ir = block["internal_risk"]
        internal_risk = InternalRiskData(
            overall_risk=ir.get("overall_risk", 0.0),
            ai_likelihood=ir.get("ai_likelihood"),
            template_score=ir.get("template_score"),
            semantic_empty_score=ir.get("semantic_empty_score"),
            repetition_score=ir.get("repetition_score"),
            citation_risk=ir.get("citation_risk"),
            reasons=ir.get("reasons", []),
        )

    return DocumentBlockResponse(
        block_id=block["block_id"],
        block_type=block["block_type"],
        text=block["text"],
        effective_text=effective_text,
        risk_score=block.get("risk_score"),
        report_risk=report_risk,
        internal_risk=internal_risk,
        rewrite_status="rewritten" if patch else "none",
        source_type=block["source_type"],
        source_map=block.get("source_map"),
        section_title=block.get("section_title"),
        section_type=block.get("section_type"),
        char_count=block.get("char_count", 0),
        display_order=block["display_order"],
    )


@router.get("/runs/{run_id}/blocks")
async def get_run_blocks(
    run_id: str,
    auth: AuthContext = Depends(get_auth_context),
) -> dict[str, Any]:
    """获取某个 run 下的所有 blocks（含最新 patch 后的 effective_text）。"""
    repository = get_repository()
    db_run = repository.get_run(run_id)
    if db_run is None:
        raise HTTPException(status_code=404, detail="run not found")
    ensure_document_access(
        document_id=str(db_run["document_id"]), auth=auth, repository=repository
    )

    blocks = repository.list_document_blocks(str(db_run["document_id"]))
    patches = repository.list_latest_patches_by_run(
        str(db_run["document_id"]), run_id
    )
    patch_map = {p["block_id"]: p for p in patches}

    result = [_build_block_response(b, patch_map.get(b["block_id"])) for b in blocks]
    reports = repository.list_cnki_reports_by_document(str(db_run["document_id"]))
    report_summary = None
    unmatched_spans: list[dict[str, Any]] = []
    if reports:
        latest_report = reports[0]
        counts = {"high": 0, "medium": 0, "low": 0}
        report_id = str(latest_report.get("id") or "")
        spans = _list_report_spans(repository, report_id) if report_id else []
        if spans:
            for span in spans:
                level = span.get("risk_level")
                if level in counts:
                    counts[level] += 1
            unmatched_spans = _list_unmatched_report_spans(
                repository,
                document_id=str(db_run["document_id"]),
                report_id=report_id,
            )
        else:
            for block in blocks:
                report_risk = block.get("report_risk")
                if isinstance(report_risk, dict):
                    level = report_risk.get("risk_level")
                    if level in counts:
                        counts[level] += 1
        report_summary = {
            "reportType": latest_report.get("report_type"),
            "totalCopyRatio": latest_report.get("total_copy_ratio"),
            "aigcRatio": latest_report.get("aigc_ratio"),
            "highRiskCount": counts["high"],
            "mediumRiskCount": counts["medium"],
            "lowRiskCount": counts["low"],
            "unmatchedCount": len(unmatched_spans),
        }

    return {
        "blocks": result,
        "reportSummary": report_summary,
        "unmatchedSpans": [_span_response(item) for item in unmatched_spans],
    }


# ------------------------------------------------------------------
# Document Patches
# ------------------------------------------------------------------

@router.post("/runs/{run_id}/patches")
async def create_patch(
    run_id: str,
    payload: DocumentPatchRequest,
    auth: AuthContext = Depends(get_auth_context),
) -> DocumentPatchResponse:
    """为某个 block 创建改写 patch。"""
    repository = get_repository()
    db_run = repository.get_run(run_id)
    if db_run is None:
        raise HTTPException(status_code=404, detail="run not found")
    ensure_document_access(
        document_id=str(db_run["document_id"]), auth=auth, repository=repository
    )

    patch = repository.insert_document_patch(
        document_id=str(db_run["document_id"]),
        run_id=run_id,
        block_id=payload.block_id,
        old_text=payload.old_text,
        new_text=payload.new_text,
        source_map=payload.source_map,
        created_by=auth.user_id if auth.authenticated else None,
    )

    return DocumentPatchResponse(
        id=str(patch["id"]),
        document_id=str(patch["document_id"]),
        run_id=str(patch["run_id"]) if patch.get("run_id") else None,
        block_id=patch["block_id"],
        old_text=patch["old_text"],
        new_text=patch["new_text"],
        created_at=patch["created_at"].isoformat() if patch.get("created_at") else "",
    )


@router.get("/runs/{run_id}/patches")
async def list_patches(
    run_id: str,
    auth: AuthContext = Depends(get_auth_context),
) -> list[DocumentPatchResponse]:
    """列出某个 run 下的所有 patches。"""
    repository = get_repository()
    db_run = repository.get_run(run_id)
    if db_run is None:
        raise HTTPException(status_code=404, detail="run not found")
    ensure_document_access(
        document_id=str(db_run["document_id"]), auth=auth, repository=repository
    )

    patches = repository.list_document_patches(str(db_run["document_id"]), run_id)
    return [
        DocumentPatchResponse(
            id=str(p["id"]),
            document_id=str(p["document_id"]),
            run_id=str(p["run_id"]) if p.get("run_id") else None,
            block_id=p["block_id"],
            old_text=p["old_text"],
            new_text=p["new_text"],
            created_at=p["created_at"].isoformat() if p.get("created_at") else "",
        )
        for p in patches
    ]
