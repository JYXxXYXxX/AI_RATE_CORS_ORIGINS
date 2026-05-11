"""Document Block 解析服务。

核心职责：把各种格式（docx/pdf/doc/txt）统一解析成 DocumentBlock 列表。
所有上传的文档都会经过这里生成 blocks，供前端编辑层使用。
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from app.services.document_loader import _decode_text
from app.services.text_processing import SECTION_HEADING_RE


@dataclass
class DocumentBlock:
    block_id: str
    block_type: str
    text: str
    html: str | None
    source_type: str
    source_map: dict | None
    paragraph_index: int | None
    section_title: str | None
    section_type: str | None
    char_count: int
    display_order: int

    def to_dict(self) -> dict[str, Any]:
        return {
            "block_id": self.block_id,
            "block_type": self.block_type,
            "text": self.text,
            "html": self.html,
            "source_type": self.source_type,
            "source_map": self.source_map or {},
            "paragraph_index": self.paragraph_index,
            "section_title": self.section_title,
            "section_type": self.section_type,
            "char_count": self.char_count,
            "display_order": self.display_order,
        }


def parse_document_to_blocks(
    file_path: str,
    source_type: str,
) -> list[DocumentBlock]:
    """统一解析入口。"""
    path = Path(file_path)
    if not path.exists():
        return []

    content = path.read_bytes()
    lower = source_type.lower()

    if lower == "docx":
        return _parse_docx_blocks(content)
    if lower == "pdf":
        return _parse_pdf_blocks(content)
    if lower == "doc":
        # doc 先尝试当作文本提取，后续阶段会加 docx 转换
        return _parse_text_blocks(content)

    # txt / md / 其他
    return _parse_text_blocks(content)


def _parse_docx_blocks(content: bytes) -> list[DocumentBlock]:
    """解析 docx 为 blocks，精确记录 paragraphIndex。"""
    blocks: list[DocumentBlock] = []
    try:
        doc = Document(BytesIO(content))
    except Exception:
        return _parse_text_blocks(content)

    current_section_title: str | None = None
    order = 0

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text:
            continue

        # 判断是否为标题
        block_type = _detect_block_type(text)
        if block_type == "heading":
            current_section_title = text.strip()

        blocks.append(
            DocumentBlock(
                block_id=f"p_{para_idx:04d}",
                block_type=block_type,
                text=text,
                html=None,
                source_type="docx",
                source_map={"paragraphIndex": para_idx},
                paragraph_index=para_idx,
                section_title=current_section_title,
                section_type=_infer_section_type(current_section_title),
                char_count=len(text),
                display_order=order,
            )
        )
        order += 1

    # 提取表格作为独立 block（简化处理：只记录文本摘要）
    for tbl_idx, table in enumerate(doc.tables):
        texts: list[str] = []
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                texts.append(" | ".join(row_texts))
        if texts:
            table_text = "\n".join(texts)
            blocks.append(
                DocumentBlock(
                    block_id=f"t_{tbl_idx:04d}",
                    block_type="table",
                    text=table_text,
                    html=None,
                    source_type="docx",
                    source_map={"tableIndex": tbl_idx},
                    paragraph_index=None,
                    section_title=current_section_title,
                    section_type=_infer_section_type(current_section_title),
                    char_count=len(table_text),
                    display_order=order,
                )
            )
            order += 1

    return blocks


def _parse_pdf_blocks(content: bytes) -> list[DocumentBlock]:
    """解析 PDF 为 blocks，按页分组。"""
    blocks: list[DocumentBlock] = []
    try:
        reader = PdfReader(BytesIO(content))
    except Exception:
        return _parse_text_blocks(content)

    order = 0
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            continue

        # 按段落拆分
        paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
        for para_idx, para_text in enumerate(paragraphs):
            block_type = _detect_block_type(para_text)
            blocks.append(
                DocumentBlock(
                    block_id=f"page{page_num}_p{para_idx:04d}",
                    block_type=block_type,
                    text=para_text,
                    html=None,
                    source_type="pdf",
                    source_map={"pageNumber": page_num, "paragraphIndex": para_idx},
                    paragraph_index=para_idx,
                    section_title=None,
                    section_type=None,
                    char_count=len(para_text),
                    display_order=order,
                )
            )
            order += 1

    return blocks


def _parse_text_blocks(content: bytes) -> list[DocumentBlock]:
    """解析纯文本为 blocks。"""
    text = _decode_text(content)
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]

    current_section_title: str | None = None
    blocks: list[DocumentBlock] = []

    for idx, para_text in enumerate(paragraphs):
        block_type = _detect_block_type(para_text)
        if block_type == "heading":
            current_section_title = para_text

        blocks.append(
            DocumentBlock(
                block_id=f"p_{idx:04d}",
                block_type=block_type,
                text=para_text,
                html=None,
                source_type="txt",
                source_map={"paragraphIndex": idx},
                paragraph_index=idx,
                section_title=current_section_title,
                section_type=_infer_section_type(current_section_title),
                char_count=len(para_text),
                display_order=idx,
            )
        )

    return blocks


def _detect_block_type(text: str) -> str:
    """判断段落类型。"""
    t = text.strip()
    if not t:
        return "unknown"

    # 标题检测
    if SECTION_HEADING_RE.match(t):
        return "heading"

    # 论文标题检测：第一行、无标点、长度适中
    if len(t) <= 80 and not re.search(r"[。！？!?,，;；]", t):
        # 进一步判断：如果是整篇文档的第一行，很可能是标题
        # 这个逻辑会在上层根据 display_order == 0 补充
        pass

    return "paragraph"


def _infer_section_type(section_title: str | None) -> str | None:
    """从章节标题推断 section_type。"""
    if not section_title:
        return None

    t = section_title.lower()
    mapping = {
        "摘要": "abstract",
        "abstract": "abstract",
        "引言": "introduction",
        "introduction": "introduction",
        "绪论": "introduction",
        "文献综述": "review",
        "相关研究": "review",
        "研究方法": "method",
        "实验方法": "method",
        "method": "method",
        "结果": "result",
        "results": "result",
        "实验结果": "result",
        "讨论": "discussion",
        "discussion": "discussion",
        "结论": "conclusion",
        "总结": "conclusion",
        "conclusion": "conclusion",
        "参考文献": "references",
        "references": "references",
        "致谢": "acknowledgement",
        "acknowledgements": "acknowledgement",
    }

    for key, val in mapping.items():
        if key in t:
            return val
    return "body"


def blocks_to_sections(
    blocks: list[DocumentBlock],
    target_chars: int = 350,
) -> list[dict[str, Any]]:
    """把 blocks 转成现有的 section 格式，供分析流程兼容使用。

    只处理 paragraph/heading 类型的 block，按 display_order 合并长段落。
    """
    from app.services.text_processing import SENTENCE_BOUNDARY_RE

    sections: list[dict[str, Any]] = []
    sec_idx = 0

    for block in blocks:
        if block.block_type not in ("paragraph", "heading"):
            continue

        text = block.text
        if block.block_type == "heading":
            # 标题单独成段，但通常很短，直接作为一个 section
            sections.append(
                {
                    "section_index": sec_idx,
                    "paragraph_index": block.paragraph_index,
                    "section_title": block.section_title,
                    "section_type": block.section_type or "body",
                    "text_preview": text[:200],
                    "content": text,
                    "char_count": len(text),
                }
            )
            sec_idx += 1
            continue

        # 普通段落：如果太长，按句子边界拆分
        if len(text) <= target_chars * 1.35:
            sections.append(
                {
                    "section_index": sec_idx,
                    "paragraph_index": block.paragraph_index,
                    "section_title": block.section_title,
                    "section_type": block.section_type or "body",
                    "text_preview": text[:200],
                    "content": text,
                    "char_count": len(text),
                }
            )
            sec_idx += 1
        else:
            sentences = SENTENCE_BOUNDARY_RE.split(text)
            buf = ""
            for sent in sentences:
                if not sent.strip():
                    continue
                if len(buf) + len(sent) < target_chars * 1.35:
                    buf += sent
                else:
                    if buf:
                        sections.append(
                            {
                                "section_index": sec_idx,
                                "paragraph_index": block.paragraph_index,
                                "section_title": block.section_title,
                                "section_type": block.section_type or "body",
                                "text_preview": buf[:200],
                                "content": buf,
                                "char_count": len(buf),
                            }
                        )
                        sec_idx += 1
                    buf = sent
            if buf:
                sections.append(
                    {
                        "section_index": sec_idx,
                        "paragraph_index": block.paragraph_index,
                        "section_title": block.section_title,
                        "section_type": block.section_type or "body",
                        "text_preview": buf[:200],
                        "content": buf,
                        "char_count": len(buf),
                    }
                )
                sec_idx += 1

    return sections
