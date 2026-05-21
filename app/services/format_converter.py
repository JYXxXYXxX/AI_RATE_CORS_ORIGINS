from __future__ import annotations

import shutil
import subprocess
import tempfile
import re
from dataclasses import dataclass
from pathlib import Path
from uuid import uuid4

from docx import Document

from app.services.document_loader import extract_text


SUPPORTED_CONVERSION_EXTENSIONS = {".doc", ".pdf", ".rtf", ".odt"}


@dataclass(frozen=True)
class ConvertedDocument:
    path: Path
    filename: str
    engine: str


def convert_to_docx(
    *,
    filename: str,
    content: bytes,
    output_dir: str,
    timeout_seconds: int = 300,
) -> ConvertedDocument:
    source_name = Path(filename or "paper").name
    suffix = Path(source_name).suffix.lower()
    if suffix not in SUPPORTED_CONVERSION_EXTENSIONS:
        raise ValueError("仅支持 .doc、.pdf、.rtf、.odt 转换为 .docx")
    if not content:
        raise ValueError("文件为空")

    out_dir = Path(output_dir) / "conversions"
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = _safe_stem(Path(source_name).stem)
    output_name = f"{stem}.docx"
    output_path = out_dir / f"{uuid4().hex}_{output_name}"

    if suffix == ".pdf":
        return _convert_pdf_to_docx(
            source_name=source_name,
            content=content,
            output_path=output_path,
            output_name=output_name,
        )
    if suffix == ".doc":
        return _convert_doc_text_to_docx(
            source_name=source_name,
            content=content,
            output_path=output_path,
            output_name=output_name,
        )

    return _convert_office_to_docx(
        source_name=source_name,
        content=content,
        output_path=output_path,
        output_name=output_name,
        timeout_seconds=timeout_seconds,
    )


def _convert_pdf_to_docx(
    *,
    source_name: str,
    content: bytes,
    output_path: Path,
    output_name: str,
) -> ConvertedDocument:
    try:
        from pdf2docx import Converter
    except ImportError as exc:
        raise RuntimeError("PDF 转 Word 组件未安装，请联系管理员配置 pdf2docx") from exc

    if not content.startswith(b"%PDF"):
        raise ValueError("文件内容不是有效的 PDF")

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(content)
        source_path = Path(tmp.name)
    try:
        converter = Converter(str(source_path))
        try:
            converter.convert(str(output_path))
        finally:
            converter.close()
    except Exception as exc:
        raise RuntimeError(
            f"{source_name} 转换失败。扫描版 PDF 可能需要先 OCR，再重新上传。"
        ) from exc
    finally:
        source_path.unlink(missing_ok=True)

    if not output_path.exists() or output_path.stat().st_size == 0:
        raise RuntimeError("PDF 转换没有生成有效的 DOCX 文件")
    return ConvertedDocument(path=output_path, filename=output_name, engine="pdf2docx")


def _convert_doc_text_to_docx(
    *,
    source_name: str,
    content: bytes,
    output_path: Path,
    output_name: str,
) -> ConvertedDocument:
    try:
        text = extract_text(source_name, content)
    except Exception as exc:
        raise RuntimeError(
            "旧版 .doc 解析失败。请先用 Word/WPS 打开并另存为 .docx 后再上传。"
        ) from exc
    text = _xml_compatible_text(text)
    if not text.strip():
        raise RuntimeError("旧版 .doc 没有提取到可转换的正文内容")

    document = Document()
    document.add_heading(Path(source_name).stem or "Converted document", level=1)
    for chunk in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        paragraph = chunk.strip()
        if paragraph:
            document.add_paragraph(paragraph)
    document.save(output_path)

    if not output_path.exists() or output_path.stat().st_size == 0:
        raise RuntimeError("旧版 .doc 转换没有生成有效的 DOCX 文件")
    return ConvertedDocument(path=output_path, filename=output_name, engine="doc-text")


def _convert_office_to_docx(
    *,
    source_name: str,
    content: bytes,
    output_path: Path,
    output_name: str,
    timeout_seconds: int,
) -> ConvertedDocument:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        raise RuntimeError("Word 格式转换组件未安装，请联系管理员配置 LibreOffice")

    suffix = Path(source_name).suffix.lower() or ".doc"
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        source_path = temp_path / source_name
        source_path.write_bytes(content)
        profile_dir = temp_path / "lo-profile"
        cmd = [
            soffice,
            "--headless",
            f"-env:UserInstallation=file:///{profile_dir.as_posix()}",
            "--convert-to",
            "docx",
            "--outdir",
            str(temp_path),
            str(source_path),
        ]
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=timeout_seconds,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError(
                "转换超时。这个文件可能较大、受保护或结构复杂，请先尝试用 Word/WPS 另存为 .docx 后再上传。"
            ) from exc
        if result.returncode != 0:
            detail = (result.stderr or result.stdout or "").strip()
            raise RuntimeError(detail or "LibreOffice 转换失败")

        converted = temp_path / f"{source_path.stem}.docx"
        if not converted.exists():
            matches = list(temp_path.glob("*.docx"))
            converted = matches[0] if matches else converted
        if not converted.exists() or converted.stat().st_size == 0:
            raise RuntimeError(f"{suffix} 转换没有生成有效的 DOCX 文件")
        shutil.copyfile(converted, output_path)

    return ConvertedDocument(path=output_path, filename=output_name, engine="libreoffice")


def _safe_stem(value: str) -> str:
    safe = "".join(ch if ch.isalnum() or ch in {"-", "_"} else "_" for ch in value)
    return safe.strip("_")[:80] or "converted"


_INVALID_XML_CHARS = re.compile(
    r"[\x00-\x08\x0b\x0c\x0e-\x1f\ufffe\uffff]"
)


def _xml_compatible_text(value: str) -> str:
    return _INVALID_XML_CHARS.sub("", value)
