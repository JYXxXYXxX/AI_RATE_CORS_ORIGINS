"""DOCX Patch 服务。

核心职责：基于原始 docx 母版，精确应用 text patches。
关键改进：通过 paragraphIndex 直接定位段落，而不是模糊字符串匹配。
"""

from __future__ import annotations

from dataclasses import dataclass, field
import io
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn


RISK_COLORS = {
    "high": "FFCDD2",
    "medium": "FFE0B2",
    "low": "E1BEE7",
    "normal": "C8E6C9",
}


@dataclass
class DocxPatchFailure:
    block_id: str
    paragraph_index: int | None
    reason: str
    old_text_preview: str


@dataclass
class DocxPatchStats:
    requested_patch_count: int
    applied_count: int = 0
    failed_count: int = 0
    skipped_block_count: int = 0
    highlighted_block_count: int = 0
    failures: list[DocxPatchFailure] = field(default_factory=list)

    def to_headers(self) -> dict[str, str]:
        return {
            "X-PataFix-Patch-Requested": str(self.requested_patch_count),
            "X-PataFix-Patch-Applied": str(self.applied_count),
            "X-PataFix-Patch-Failed": str(self.failed_count),
            "X-PataFix-Patch-Skipped": str(self.skipped_block_count),
            "X-PataFix-Patch-Highlighted": str(self.highlighted_block_count),
        }


@dataclass
class DocxPatchResult:
    content: bytes
    stats: DocxPatchStats


class DocxPatchError(ValueError):
    def __init__(self, stats: DocxPatchStats) -> None:
        self.stats = stats
        details = "; ".join(
            f"{item.block_id}: {item.reason}" for item in stats.failures[:5]
        )
        super().__init__(details or "DOCX patch failed")


def _set_para_shading(paragraph, color_hex: str) -> None:
    pPr = paragraph._element.get_or_add_pPr()
    shading = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:fill="{color_hex}" w:val="clear"/>'
    )
    pPr.append(shading)


import re


def _try_replace_in_runs(paragraph, old_text: str, new_text: str) -> bool:
    """尝试在段落内跨 run 替换文本，保留格式。

    策略：
    1. 收集所有 run 的文本，拼接成完整字符串
    2. 找到 old_text 的位置（支持忽略空白差异的模糊匹配）
    3. 计算跨越了哪些 run
    4. 第一个 run 保留样式并承载 prefix + new_text
    5. 中间 runs 清空
    6. 最后一个 run 承载 suffix
    """
    runs = paragraph.runs
    if not runs:
        return False

    # 收集所有文本和位置映射
    run_texts = []
    total = ""
    for run in runs:
        t = run.text
        run_texts.append(t)
        total += t

    # 查找 old_text
    idx = total.find(old_text)
    match_len = len(old_text)
    if idx == -1:
        # 尝试忽略空白差异的模糊匹配
        parts = [part for part in old_text.split() if part]
        if not parts:
            return False
        pattern = r"\s*".join(re.escape(part) for part in parts)
        m = re.search(pattern, total)
        if m:
            idx = m.start()
            match_len = m.end() - m.start()
        else:
            return False

    end_idx = idx + match_len

    # 计算 old_text 跨越了哪些 run
    current_pos = 0
    start_run = -1
    end_run = -1
    for i, t in enumerate(run_texts):
        run_start = current_pos
        run_end = current_pos + len(t)
        if start_run == -1 and run_end > idx:
            start_run = i
        if start_run != -1 and run_end >= end_idx:
            end_run = i
            break
        current_pos = run_end

    if start_run == -1 or end_run == -1:
        return False

    # 计算在 start_run 和 end_run 中的偏移
    prefix_len = idx - sum(len(run_texts[j]) for j in range(start_run))
    suffix_len = sum(len(run_texts[j]) for j in range(end_run + 1)) - end_idx

    prefix = run_texts[start_run][:prefix_len]
    suffix = run_texts[end_run][len(run_texts[end_run]) - suffix_len:]

    # 替换
    runs[start_run].text = prefix + new_text + suffix
    for i in range(start_run + 1, end_run + 1):
        runs[i].text = ""

    return True


def _fallback_replace_paragraph(paragraph, new_text: str) -> None:
    """兜底方案：清空所有 runs，用第一个 run 的样式写新文本。"""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def _source_map_replacements(patch: dict) -> list[dict[str, str]]:
    source_map = patch.get("source_map") or {}
    raw = source_map.get("replacements") if isinstance(source_map, dict) else None
    if not isinstance(raw, list):
        return []
    replacements: list[dict[str, str]] = []
    for item in raw:
        if not isinstance(item, dict):
            continue
        old_text = str(item.get("old_text") or "").strip()
        new_text = str(item.get("new_text") or "").strip()
        if old_text and new_text and old_text != new_text:
            replacements.append({"old_text": old_text, "new_text": new_text})
    return replacements


def export_docx_with_patches(
    original_path: str,
    blocks: list[dict],
    patches: list[dict],
    *,
    highlight_risks: bool = False,
    strict: bool = False,
) -> bytes:
    """基于原始 docx 母版，仅应用文本 patches。

    默认不写入风险底色，避免导出的改写稿破坏用户原论文格式。风险颜色只用于
    网页编辑器和专门的高亮导出视图；正式改写稿走“仅文本替换”路线。

    Args:
        original_path: 原始 docx 文件路径
        blocks: 文档 blocks 列表，每个 block 包含 block_id, text, risk_score, source_map
        patches: 改写 patches 列表，每个 patch 包含 block_id, old_text, new_text
        highlight_risks: 是否在导出 docx 中写入风险底色。

    Returns:
        docx 文件的字节内容
    """
    return export_docx_with_patch_report(
        original_path,
        blocks,
        patches,
        highlight_risks=highlight_risks,
        strict=strict,
    ).content


def export_docx_with_patch_report(
    original_path: str,
    blocks: list[dict],
    patches: list[dict],
    *,
    highlight_risks: bool = False,
    strict: bool = False,
) -> DocxPatchResult:
    path = Path(original_path)
    if not path.exists() or path.suffix.lower() != ".docx":
        raise ValueError("原始 docx 文件不存在")

    doc = Document(original_path)
    patch_map = {p["block_id"]: p for p in patches}
    handled_patch_ids: set[str] = set()
    stats = DocxPatchStats(requested_patch_count=len(patches))

    for block in blocks:
        block_id = block.get("block_id")
        if block.get("block_type") != "paragraph":
            if block_id in patch_map:
                handled_patch_ids.add(str(block_id))
                stats.skipped_block_count += 1
                stats.failures.append(
                    DocxPatchFailure(
                        block_id=str(block_id),
                        paragraph_index=None,
                        reason="该 block 不是正文段落，已跳过",
                        old_text_preview=str(patch_map[block_id].get("old_text", ""))[:80],
                    )
                )
            continue

        patch = patch_map.get(block_id)
        para_idx = block.get("source_map", {}).get("paragraphIndex")
        if para_idx is None or para_idx >= len(doc.paragraphs):
            if patch:
                handled_patch_ids.add(str(block_id))
                stats.failed_count += 1
                stats.failures.append(
                    DocxPatchFailure(
                        block_id=str(block_id),
                        paragraph_index=para_idx,
                        reason="无法定位原文档段落",
                        old_text_preview=str(patch.get("old_text", ""))[:80],
                    )
                )
            continue

        para = doc.paragraphs[para_idx]

        if patch:
            handled_patch_ids.add(str(block_id))
            old_text = patch["old_text"]
            new_text = patch["new_text"]
            replacements = _source_map_replacements(patch)
            if replacements:
                failed = [
                    item for item in replacements
                    if not _try_replace_in_runs(para, item["old_text"], item["new_text"])
                ]
                if not failed:
                    stats.applied_count += 1
                elif strict:
                    stats.failed_count += 1
                    stats.failures.append(
                        DocxPatchFailure(
                            block_id=str(block_id),
                            paragraph_index=para_idx,
                            reason="部分逐句替换未在对应段落中找到，未执行整段覆盖以避免破坏格式",
                            old_text_preview=str(failed[0].get("old_text", ""))[:80],
                        )
                    )
                elif _try_replace_in_runs(para, old_text, new_text):
                    stats.applied_count += 1
                else:
                    _fallback_replace_paragraph(para, new_text)
                    stats.applied_count += 1
            elif _try_replace_in_runs(para, old_text, new_text):
                stats.applied_count += 1
            elif strict:
                stats.failed_count += 1
                stats.failures.append(
                    DocxPatchFailure(
                        block_id=str(block_id),
                        paragraph_index=para_idx,
                        reason="原句未在对应段落中找到，未执行整段覆盖以避免破坏格式",
                        old_text_preview=str(old_text)[:80],
                    )
                )
            else:
                _fallback_replace_paragraph(para, new_text)
                stats.applied_count += 1

        if highlight_risks:
            color = RISK_COLORS["normal"]
            report_risk = block.get("report_risk")
            if isinstance(report_risk, dict) and report_risk.get("risk_level"):
                color = RISK_COLORS.get(report_risk.get("risk_level"), RISK_COLORS["normal"])
            else:
                risk_score = block.get("risk_score") or 0
                if risk_score >= 70:
                    color = RISK_COLORS["high"]
                elif risk_score >= 60:
                    color = RISK_COLORS["medium"]
                elif risk_score >= 30:
                    color = RISK_COLORS["low"]
            _set_para_shading(para, color)
            stats.highlighted_block_count += 1

    missing_patch_ids = set(str(item.get("block_id")) for item in patches) - handled_patch_ids
    for block_id in sorted(missing_patch_ids):
        patch = patch_map.get(block_id)
        if patch is None:
            continue
        stats.failed_count += 1
        stats.failures.append(
            DocxPatchFailure(
                block_id=block_id,
                paragraph_index=None,
                reason="patch 未匹配到任何可导出的正文 block",
                old_text_preview=str(patch.get("old_text", ""))[:80],
            )
        )

    if strict and stats.failed_count:
        raise DocxPatchError(stats)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return DocxPatchResult(content=buf.getvalue(), stats=stats)
