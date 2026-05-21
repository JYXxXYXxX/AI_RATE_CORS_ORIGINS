from __future__ import annotations

from io import BytesIO

from docx import Document

from app.services.docx_patch import (
    DocxPatchError,
    export_docx_with_patch_report,
    export_docx_with_patches,
)


def test_docx_patch_replaces_text_without_risk_shading(tmp_path) -> None:
    original = tmp_path / "original.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("本文通过")
    paragraph.add_run("系统设计")
    paragraph.add_run("提升财务管理效率。")
    doc.save(original)

    blocks = [
        {
            "block_id": "b1",
            "block_type": "paragraph",
            "text": "本文通过系统设计提升财务管理效率。",
            "source_map": {"paragraphIndex": 0},
            "risk_score": 82,
        }
    ]
    patches = [
        {
            "block_id": "b1",
            "old_text": "本文通过系统设计提升财务管理效率。",
            "new_text": "本系统在预算提醒模块中读取收支记录，并据此生成财务提醒。",
        }
    ]

    output = export_docx_with_patches(str(original), blocks, patches)
    patched = Document(BytesIO(output))

    assert patched.paragraphs[0].text == "本系统在预算提醒模块中读取收支记录，并据此生成财务提醒。"
    p_pr_xml = patched.paragraphs[0]._element.get_or_add_pPr().xml
    assert "<w:shd" not in p_pr_xml

    report = export_docx_with_patch_report(str(original), blocks, patches, strict=True)
    assert report.stats.requested_patch_count == 1
    assert report.stats.applied_count == 1
    assert report.stats.failed_count == 0


def test_docx_patch_can_optionally_export_highlighted_copy(tmp_path) -> None:
    original = tmp_path / "original.docx"
    doc = Document()
    doc.add_paragraph("这是一个高风险段落。")
    doc.save(original)

    output = export_docx_with_patches(
        str(original),
        [
            {
                "block_id": "b1",
                "block_type": "paragraph",
                "text": "这是一个高风险段落。",
                "source_map": {"paragraphIndex": 0},
                "risk_score": 82,
            }
        ],
        [],
        highlight_risks=True,
    )
    patched = Document(BytesIO(output))

    assert "<w:shd" in patched.paragraphs[0]._element.get_or_add_pPr().xml


def test_docx_patch_applies_sentence_replacements_without_paragraph_overwrite(tmp_path) -> None:
    original = tmp_path / "original.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("原文第一句。")
    bold_run = paragraph.add_run("需要逐字修改的句子。")
    bold_run.bold = True
    paragraph.add_run("原文第三句。")
    doc.save(original)

    blocks = [
        {
            "block_id": "b1",
            "block_type": "paragraph",
            "text": "原文第一句。需要逐字修改的句子。原文第三句。",
            "source_map": {"paragraphIndex": 0},
        }
    ]
    patches = [
        {
            "block_id": "b1",
            "old_text": "原文第一句。需要逐字修改的句子。原文第三句。",
            "new_text": "原文第一句。需要局部调整的句子。原文第三句。",
            "source_map": {
                "replacements": [
                    {
                        "old_text": "需要逐字修改的句子。",
                        "new_text": "需要局部调整的句子。",
                    }
                ]
            },
        }
    ]

    report = export_docx_with_patch_report(str(original), blocks, patches, strict=True)
    patched = Document(BytesIO(report.content))

    assert patched.paragraphs[0].text == "原文第一句。需要局部调整的句子。原文第三句。"
    assert patched.paragraphs[0].runs[1].bold is True
    assert report.stats.applied_count == 1
    assert report.stats.failed_count == 0


def test_docx_patch_strict_mode_refuses_unsafe_paragraph_overwrite(tmp_path) -> None:
    original = tmp_path / "original.docx"
    doc = Document()
    doc.add_paragraph("原文段落中没有待替换的完整旧句。")
    doc.save(original)

    blocks = [
        {
            "block_id": "b1",
            "block_type": "paragraph",
            "text": "原文段落中没有待替换的完整旧句。",
            "source_map": {"paragraphIndex": 0},
        }
    ]
    patches = [
        {
            "block_id": "b1",
            "old_text": "这句在原文档中不存在。",
            "new_text": "新句子不应该通过整段覆盖写进去。",
        }
    ]

    try:
        export_docx_with_patch_report(str(original), blocks, patches, strict=True)
    except DocxPatchError as exc:
        assert exc.stats.requested_patch_count == 1
        assert exc.stats.applied_count == 0
        assert exc.stats.failed_count == 1
        assert "未执行整段覆盖" in exc.stats.failures[0].reason
    else:
        raise AssertionError("strict patching should reject unsafe overwrite fallback")

    non_strict = export_docx_with_patch_report(
        str(original), blocks, patches, strict=False
    )
    patched = Document(BytesIO(non_strict.content))
    assert patched.paragraphs[0].text == "新句子不应该通过整段覆盖写进去。"
